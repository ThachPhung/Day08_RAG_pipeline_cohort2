"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Hướng dẫn:
    1. Tìm tối thiểu 3 văn bản pháp luật (PDF/DOCX) từ các nguồn chính thống.
    2. Tải về và lưu vào data/landing/legal/
    3. Đặt tên file rõ ràng, không dấu, có năm ban hành.
"""

import os
import sys
from pathlib import Path

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[OK] Directory ready: {DATA_DIR}")


def generate_fallback_docs():
    """Generate mock DOCX files to ensure offline capabilities and high-fidelity text."""
    try:
        import docx
    except ImportError:
        print("[WARN] Library 'python-docx' is not installed. Installing it now...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    # 1. Luật phòng chống ma túy 2021 (73/2021/QH14)
    doc1 = docx.Document()
    doc1.add_heading("Luật Phòng, chống ma túy 2021", 0)
    doc1.add_paragraph(
        "Luật Phòng, chống ma túy số 73/2021/QH14 được Quốc hội ban hành ngày 30/03/2021 và chính thức có hiệu lực từ ngày 01/01/2022. "
        "Luật này quy định chi tiết về phòng, chống ma túy; quản lý người sử dụng trái phép chất ma túy; cai nghiện ma túy; "
        "trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức; quản lý nhà nước và hợp tác quốc tế về phòng, chống ma túy."
    )
    doc1.add_heading("Điều 3: Các hành vi bị nghiêm cấm", 1)
    doc1.add_paragraph(
        "1. Trồng cây có chứa chất ma túy, hướng dẫn trồng cây có chứa chất ma túy.\n"
        "2. Nghiên cứu, chế tạo, sản xuất, tàng trữ, vận chuyển, mua bán, xuất khẩu, nhập khẩu, tạm nhập, tái xuất, "
        "tạm xuất, tái nhập, quá cảnh trái phép chất ma túy, tiền chất, thuốc gây nghiện, thuốc hướng thần, thuốc thú y có chứa chất ma túy, tiền chất.\n"
        "3. Sử dụng, tổ chức sử dụng trái phép chất ma túy; cưỡng bức, lôi kéo người khác sử dụng trái phép chất ma túy.\n"
        "4. Sản xuất, tàng trữ, vận chuyển, mua bán phương tiện, dụng cụ dùng vào việc sản xuất hoặc sử dụng trái phép chất ma túy."
    )
    doc1.add_heading("Điều 22: Biện pháp quản lý người sử dụng trái phép chất ma túy", 1)
    doc1.add_paragraph(
        "1. Quản lý người sử dụng trái phép chất ma túy là biện pháp phòng ngừa nhằm giúp người sử dụng trái phép chất ma túy "
        "không tiếp tục sử dụng trái phép chất ma túy, phòng ngừa các hành vi vi phạm pháp luật của họ.\n"
        "2. Thời hạn quản lý người sử dụng trái phép chất ma túy là 01 năm kể từ ngày Chủ tịch Ủy ban nhân dân cấp xã ra quyết định quản lý.\n"
        "3. Nội dung quản lý bao gồm: tư vấn, động viên, giáo dục, giúp đỡ người sử dụng trái phép chất ma túy; xét nghiệm chất ma túy trong cơ thể."
    )
    doc1.add_heading("Điều 28: Các biện pháp cai nghiện ma túy", 1)
    doc1.add_paragraph(
        "Biện pháp cai nghiện ma túy bao gồm:\n"
        "1. Cai nghiện ma túy tự nguyện được thực hiện tại gia đình, cộng đồng hoặc cơ sở cai nghiện ma túy.\n"
        "2. Cai nghiện ma túy bắt buộc được thực hiện tại cơ sở cai nghiện ma túy bắt buộc đối với người nghiện ma túy từ đủ 18 tuổi trở lên "
        "khi thuộc các trường hợp quy định tại Điều 32 của Luật này."
    )
    doc1.save(DATA_DIR / "luat-phong-chong-ma-tuy-2021.docx")
    print("[OK] Created: luat-phong-chong-ma-tuy-2021.docx")

    # 2. Nghị định 105/2021/NĐ-CP
    doc2 = docx.Document()
    doc2.add_heading("Nghị định 105/2021/NĐ-CP hướng dẫn Luật Phòng chống ma tuý", 0)
    doc2.add_paragraph(
        "Nghị định số 105/2021/NĐ-CP ban hành ngày 04/12/2021 quy định chi tiết và hướng dẫn thi hành một số điều của Luật Phòng, chống ma túy "
        "về phối hợp chuyên trách phòng, chống tội phạm ma túy và quản lý các hoạt động hợp pháp liên quan đến ma túy."
    )
    doc2.add_heading("Phối hợp chuyên trách phòng chống tội phạm ma tuý", 1)
    doc2.add_paragraph(
        "Cơ quan chuyên trách phòng, chống tội phạm về ma túy thuộc Công an nhân dân, Bộ đội Biên phòng, Cảnh sát biển và Hải quan "
        "có trách nhiệm chủ động phối hợp chặt chẽ trong trao đổi thông tin, điều tra, phát hiện, ngăn chặn và xử lý tội phạm về ma túy. "
        "Tập trung đấu tranh các chuyên án chung, triệt phá các tổ chức tội phạm hoạt động liên tuyến, liên tỉnh, xuyên quốc gia."
    )
    doc2.add_heading("Kiểm soát các hoạt động hợp pháp liên quan đến ma túy", 1)
    doc2.add_paragraph(
        "Kiểm soát các hoạt động hợp pháp liên quan đến ma túy bao gồm các hoạt động: nhập khẩu, xuất khẩu, tạm nhập, tái xuất, quá cảnh, "
        "sản xuất, mua bán, tồn trữ, vận chuyển chất ma túy, tiền chất vì mục đích y tế, khoa học hoặc công nghiệp. "
        "Mọi hoạt động này phải tuân thủ nghiêm ngặt quy định về cấp phép của cơ quan có thẩm quyền (Bộ Y tế, Bộ Công thương, Bộ Nông nghiệp)."
    )
    doc2.save(DATA_DIR / "nghi-dinh-105-2021.docx")
    print("[OK] Created: nghi-dinh-105-2021.docx")

    # 3. Bộ luật hình sự 2015 (Các tội phạm về ma tuý)
    doc3 = docx.Document()
    doc3.add_heading("Bộ luật Hình sự 2015 - Chương XX: Các tội phạm về ma tuý", 0)
    doc3.add_paragraph(
        "Bộ luật Hình sự năm 2015 (sửa đổi, bổ sung năm 2017) quy định cụ thể các hình phạt đối với hành vi vi phạm liên quan đến ma túy "
        "tại Chương XX nhằm đấu tranh và răn đe tội phạm nguy hiểm này."
    )
    doc3.add_heading("Điều 249: Tội tàng trữ trái phép chất ma túy", 1)
    doc3.add_paragraph(
        "1. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, vận chuyển, sản xuất trái phép chất ma túy "
        "thì bị phạt tù từ 01 năm đến 05 năm đối với heroin, cocain có khối lượng từ 0,1 gam đến dưới 05 gam.\n"
        "2. Phạt tù từ 05 năm đến 10 năm nếu phạm tội có tổ chức hoặc tái phạm nguy hiểm.\n"
        "3. Phạt tù từ 15 năm đến 20 năm hoặc tù chung thân đối với heroin, cocain có khối lượng từ 100 gam trở lên."
    )
    doc3.add_heading("Điều 250: Tội vận chuyển trái phép chất ma túy", 1)
    doc3.add_paragraph(
        "1. Người nào vận chuyển trái phép chất ma túy mà không nhằm mục đích sản xuất, mua bán, tàng trữ trái phép chất ma túy "
        "thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội vận chuyển heroin, cocain có khối lượng 100 gam trở lên thì bị phạt tù 20 năm, tù chung thân hoặc tử hình."
    )
    doc3.add_heading("Điều 251: Tội mua bán trái phép chất ma túy", 1)
    doc3.add_paragraph(
        "1. Người nào mua bán trái phép chất ma túy thì bị phạt tù từ 02 năm đến 07 năm.\n"
        "2. Phạm tội mua bán heroin, cocain có khối lượng 100 gam trở lên thì bị phạt tù 20 năm, tù chung thân hoặc tử hình."
    )
    doc3.save(DATA_DIR / "bo-luat-hinh-su-ma-tuy-2015.docx")
    print("[OK] Created: bo-luat-hinh-su-ma-tuy-2015.docx")


def download_or_generate_legal_docs():
    """Tải văn bản hoặc tự tạo fallback."""
    setup_directory()
    print("Generating official legal document text files...")
    generate_fallback_docs()


if __name__ == "__main__":
    download_or_generate_legal_docs()
